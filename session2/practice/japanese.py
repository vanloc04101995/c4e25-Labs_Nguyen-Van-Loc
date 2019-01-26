from urllib.request import urlopen
from bs4 import BeautifulSoup
import pydoc
from docx import Document
from collections import OrderedDict
document = Document()

url = "https://www3.nhk.or.jp/news/html/20190121/k10011785331000.html"
client = urlopen(url)
raw_data = client.read()
content = raw_data.decode("utf8")
soup = BeautifulSoup(content,"html.parser")
# div1 = soup.find("div",id="news_textbody")
div2 = soup.find("div",id="news_textmore")

print(div2.string)
# document.add_paragraph(div1.string)
# document.add_paragraph(div2.string)
# document.save('demo.docx')